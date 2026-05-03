"""
================================================================================
multimodal_output.py — Multimodal Output Engine for JARVIS
================================================================================
Generates multimodal outputs from a single request.

Combines: text + image + voice + document + code + diagram
All processing is 100% local with optional external integrations.

Usage:
    from runtime.agency.multimodal_output import get_multimodal_engine, MultimodalOutput
    engine = get_multimodal_engine()
    output = engine.explain(
        topic="How neural networks work",
        modalities=["text", "image", "audio", "document", "diagram"],
        language="he"
    )
    print(output.text)
    print(output.image_path)
    print(output.audio_path)
================================================================================
"""

from __future__ import annotations

import os
import io
import json
import base64
import hashlib
import datetime
import textwrap
import subprocess
import tempfile
import warnings
from pathlib import Path
from dataclasses import dataclass, field, asdict
from typing import List, Dict, Optional, Any, Tuple

# ──────────────────────────────────────────────────────────────────────────────
# Optional dependency handling — graceful degradation
# ──────────────────────────────────────────────────────────────────────────────

# Matplotlib — diagram generation
try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
    _HAS_MATPLOTLIB = True
except ImportError:  # pragma: no cover
    _HAS_MATPLOTLIB = False
    warnings.warn("matplotlib not available; diagram generation will be limited.")

# NetworkX — mind map support
try:
    import networkx as nx
    _HAS_NETWORKX = True
except ImportError:  # pragma: no cover
    _HAS_NETWORKX = False

# NumPy — diagram math
try:
    import numpy as np
    _HAS_NUMPY = True
except ImportError:  # pragma: no cover
    _HAS_NUMPY = False

# PIL — image processing
try:
    from PIL import Image, ImageDraw, ImageFont
    _HAS_PIL = True
except ImportError:  # pragma: no cover
    _HAS_PIL = False

# TTS / gTTS — text-to-speech
try:
    from gtts import gTTS
    _HAS_GTTS = True
except ImportError:  # pragma: no cover
    _HAS_GTTS = False

# pyttsx3 — offline TTS
try:
    import pyttsx3
    _HAS_PYTTSX3 = True
except ImportError:  # pragma: no cover
    _HAS_PYTTSX3 = False

# ReportLab — PDF generation
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        Image as RLImage, PageBreak, KeepTogether
    )
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
    _HAS_REPORTLAB = True
except ImportError:  # pragma: no cover
    _HAS_REPORTLAB = False

# python-docx — Word document generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _HAS_DOCX = True
except ImportError:  # pragma: no cover
    _HAS_DOCX = False

# ──────────────────────────────────────────────────────────────────────────────
# Constants & Configuration
# ──────────────────────────────────────────────────────────────────────────────

DEFAULT_OUTPUT_DIR = Path.home() / ".jarvis" / "multimodal_output"
DIAGRAM_TYPES = {"flowchart", "mindmap", "timeline", "architecture", "comparison"}
DOCUMENT_TYPES = {"pdf", "docx", "html", "md"}
MODALITY_ALL = {"text", "image", "audio", "document", "code", "diagram"}

HEBREW_ALPHABET = set("אבגדהוזחטיכלמנסעפצקרשתםןץףך")
RTL_LANGUAGES = {"he", "ar", "fa", "ur"}

FONT_PATHS = [
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/System/Library/Fonts/Helvetica.ttc",
    "/Windows/Fonts/arial.ttf",
    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
]

# ──────────────────────────────────────────────────────────────────────────────
# Knowledge base: pre-built explanations for common topics
# ──────────────────────────────────────────────────────────────────────────────

KNOWLEDGE_BASE: Dict[str, Dict[str, str]] = {
    "neural network": {
        "he": (
            "רשתות נוירונים הן מערכות חישוביות בהשראת המוח האנושי. "
            "הן בנויות מיחידות עיבוד שנקראות נוירונים המחוברות ביניהן בקשרים. "
            "כל נוירון מקבל קלט, מבצע עליו חישוב, ומעביר אות לנוירונים אחרים. "
            "הרשת לומדת על ידי התאמת המשקולות בין הנוירונים כדי למזער שגיאות."
        ),
        "en": (
            "Neural networks are brain-inspired computing systems. They consist of "
            "interconnected units called neurons that process information. Each neuron "
            "receives input, performs a computation, and passes signals to other neurons. "
            "The network learns by adjusting the weights between neurons to minimize errors."
        ),
    },
    "machine learning": {
        "he": (
            "למידת מכונה היא תחום בבינה מלאכותית שבו מחשבים לומדים מנתונים. "
            "במקום לתכנת כללים מפורשים, המערכת מזהה דפוסים בנתוני אימון ומשתמשת בהם "
            "כדי לבצע חיזויים או החלטות על נתונים חדשים."
        ),
        "en": (
            "Machine Learning is a field of AI where computers learn from data. "
            "Instead of explicit programming, the system identifies patterns in training "
            "data and uses them to make predictions or decisions on new data."
        ),
    },
    "deep learning": {
        "he": (
            "למידה עמוקה היא תת-תחום של למידת מכונה המשתמש ברשתות נוירונים עם שכבות רבות. "
            "העומק של הרשת מאפשר לה ללמוד ייצוגים מורכבים ואבסטרקטיים של הנתונים. "
            "זהו הכוח שמאחורי פריצות הדרך האחרונות בראייה ממוחשבת ובעיבוד שפה טבעית."
        ),
        "en": (
            "Deep Learning is a subfield of Machine Learning using multi-layer neural networks. "
            "The depth allows the network to learn complex, hierarchical representations of data. "
            "It powers breakthroughs in computer vision and natural language processing."
        ),
    },
    "reinforcement learning": {
        "he": (
            "למידת חיזוק היא סוג של למידת מכונה שבה סוכן לומד לקבל החלטות על ידי "
            "אינטראקציה עם סביבה. הסוכן מקבל פרסים או עונשים על פעולותיו ולומד "
            "מדיניות שממקסמת את הפרס לאורך זמן."
        ),
        "en": (
            "Reinforcement Learning is a type of ML where an agent learns to make decisions "
            "by interacting with an environment. The agent receives rewards or penalties "
            "for its actions and learns a policy that maximizes cumulative reward over time."
        ),
    },
    "transformer": {
        "he": (
            "הטרנספורמר הוא ארכיטקטורת רשת נוירונים המבוססת על מנגנון תשומת הלב (Attention). "
            "הוא מאפשר למודל להתמקד בחלקים רלוונטיים של הקלט בעת עיבוד כל רכיב. "
            "הטרנספורמר מהווה את הבסיס למודלי שפה כמו GPT ו-BERT."
        ),
        "en": (
            "The Transformer is a neural network architecture based on the Attention mechanism. "
            "It allows the model to focus on relevant parts of the input when processing each element. "
            "It underpins language models like GPT and BERT."
        ),
    },
    "convolutional neural network": {
        "he": (
            "רשתות קונבולוציה (CNN) הן סוג של רשת נוירונים המיועדות לעיבוד נתונים בעלי מבנה רשתי "
            "כמו תמונות. הן משתמשות בשכבות קונבולוציה שמחלצות תכונות מקומיות כמו קצוות ומרקמים. "
            "CNNs הן הסטנדרט ברוב יישומי הראייה הממוחשבת."
        ),
        "en": (
            "Convolutional Neural Networks (CNNs) are a type of neural network designed for "
            "grid-structured data like images. They use convolutional layers that extract "
            "local features such as edges and textures. CNNs are the standard in most computer "
            "vision applications."
        ),
    },
    "recurrent neural network": {
        "he": (
            "רשתות נוירונים רקורסיביות (RNN) הן ארכיטקטורה לעיבוד נתונים סדרתיים כמו טקסט או מספרים זמניים. "
            "הן מכילות מצב פנימי שמאפשר זיכרון של קלטים קודמים. LSTM ו-GRU הן וריאנטים "
            "פופולריים שפותרים את בעיית התלות לטווח הארוך."
        ),
        "en": (
            "Recurrent Neural Networks (RNNs) process sequential data like text or time series. "
            "They maintain an internal state that allows memory of previous inputs. LSTM and GRU "
            "are popular variants that address the long-term dependency problem."
        ),
    },
    "backpropagation": {
        "he": (
            "האלגוריתם Backpropagation הוא שיטת אימון רשתות נוירונים. "
            "הוא מחשב את גרדיאנט השגיאה ביחס למשקולות על ידי הפצת השגיאה אחורנית מהשכבה האחרונה. "
            "לאחר מכן משתמשים בירידת גרדיאנט כדי לעדכן את המשקולות ולמזער את השגיאה."
        ),
        "en": (
            "Backpropagation is a training algorithm for neural networks. It computes the gradient "
            "of the error with respect to the weights by propagating the error backward from the "
            "output layer. Gradient descent then updates the weights to minimize the error."
        ),
    },
    "gradient descent": {
        "he": (
            "ירידת גרדיאנט היא אלגוריתם אופטימיזציה למציאת מינימום של פונקציה. "
            "היא עובדת על ידי הזזה בכיוון הנגדי לגרדיאנט (כיוון השיפוע התלול ביותר). "
            "בלמידת מכונה, היא משמשת לעדכון המשקולות כדי למזער את פונקציית האובדן."
        ),
        "en": (
            "Gradient Descent is an optimization algorithm for finding function minima. "
            "It iteratively moves in the direction opposite to the gradient (steepest descent). "
            "In ML, it is used to update weights to minimize the loss function."
        ),
    },
    "attention mechanism": {
        "he": (
            "מנגנון תשומת הלב מאפשר למודל להתמקד בחלקים רלוונטיים של הקלט בעת עיבוד כל רכיב. "
            "במקום לעבד את כל הקלט באופן שווה, המודל מחשב מידת הרלוונטיות של כל קלט אחר "
            "לרכיב הנוכחי ומשקלל אותם בהתאם."
        ),
        "en": (
            "The Attention mechanism lets a model focus on relevant input parts when processing "
            "each element. Instead of processing all input equally, the model computes relevance "
            "scores and weights inputs accordingly."
        ),
    },
    "generative adversarial network": {
        "he": (
            "רשתות יוצרות מתחרות (GAN) מורכבות משתי רשתות: יוצר (Generator) ומבחין (Discriminator). "
            "היוצר מייצר נתונים מזויפים והמבחין מנסה להבדיל בין נתונים אמיתיים למזויפים. "
            "התחרות ביניהן מובילה לייצור נתונים ריאליסטיים מאוד."
        ),
        "en": (
            "Generative Adversarial Networks (GANs) consist of two networks: a Generator and "
            "a Discriminator. The generator creates fake data while the discriminator tries to "
            "distinguish real from fake. Their competition produces highly realistic outputs."
        ),
    },
    "natural language processing": {
        "he": (
            "עיבוד שפה טבעית (NLP) הוא תחום בבינה מלאכותית העוסק באינטראקציה בין מחשבים לשפה אנושית. "
            "הוא כולל משימות כמו תרגום מכונה, ניתוח סנטימנט, זיהוי יישויות, ותשובות לשאלות. "
            "מודלי שפה גדולים (LLMs) מהווים את חזית התחום כיום."
        ),
        "en": (
            "Natural Language Processing (NLP) is an AI field dealing with computer-human language "
            "interaction. It includes tasks like machine translation, sentiment analysis, named "
            "entity recognition, and question answering. Large Language Models (LLMs) represent "
            "the cutting edge of the field."
        ),
    },
    "computer vision": {
        "he": (
            "ראייה ממוחשבת היא תחום בבינה מלאכותית המאפשר למחשבים להבין תוכן ויזואלי. "
            "היא כוללת זיהוי עצמים, סגמנטציה של תמונות, זיהוי פנים, וניתוח וידאו. "
            "רשתות קונבולוציה (CNNs) הן הכלי העיקרי בתחום זה."
        ),
        "en": (
            "Computer Vision is an AI field that enables computers to understand visual content. "
            "It includes object detection, image segmentation, face recognition, and video analysis. "
            "Convolutional Neural Networks (CNNs) are the primary tool in this field."
        ),
    },
}

# Code examples database
CODE_EXAMPLES: Dict[str, Dict[str, str]] = {
    "neural network": {
        "python": '''
import torch
import torch.nn as nn

class NeuralNetwork(nn.Module):
    def __init__(self, input_size, hidden_size, num_classes):
        super(NeuralNetwork, self).__init__()
        self.layer1 = nn.Linear(input_size, hidden_size)
        self.relu = nn.ReLU()
        self.dropout = nn.Dropout(0.2)
        self.layer2 = nn.Linear(hidden_size, hidden_size)
        self.layer3 = nn.Linear(hidden_size, num_classes)

    def forward(self, x):
        x = self.layer1(x)
        x = self.relu(x)
        x = self.dropout(x)
        x = self.layer2(x)
        x = self.relu(x)
        x = self.layer3(x)
        return x

# Usage
model = NeuralNetwork(input_size=784, hidden_size=256, num_classes=10)
criterion = nn.CrossEntropyLoss()
optimizer = torch.optim.Adam(model.parameters(), lr=0.001)
''',
    },
    "machine learning": {
        "python": '''
from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import train_test_split
from sklearn.metrics import accuracy_score
import pandas as pd

# Load data
data = pd.read_csv('data.csv')
X = data.drop('target', axis=1)
y = data['target']

# Split
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42
)

# Train
model = RandomForestClassifier(n_estimators=100, random_state=42)
model.fit(X_train, y_train)

# Evaluate
predictions = model.predict(X_test)
print(f"Accuracy: {accuracy_score(y_test, predictions):.2f}")
''',
    },
    "deep learning": {
        "python": '''
import torch
import torch.nn as nn

class DeepNetwork(nn.Module):
    def __init__(self):
        super().__init__()
        self.features = nn.Sequential(
            nn.Conv2d(1, 32, kernel_size=3, padding=1),
            nn.ReLU(),
            nn.MaxPool2d(2),
            nn.Conv2d(32, 64, kernel_size=3, padding=1),
            nn.ReLU(),
            nn.MaxPool2d(2),
        )
        self.classifier = nn.Sequential(
            nn.Flatten(),
            nn.Linear(64 * 7 * 7, 128),
            nn.ReLU(),
            nn.Dropout(0.5),
            nn.Linear(128, 10),
        )

    def forward(self, x):
        x = self.features(x)
        x = self.classifier(x)
        return x
''',
    },
    "transformer": {
        "python": '''
import torch
import torch.nn as nn
import math

class SelfAttention(nn.Module):
    def __init__(self, embed_size, heads):
        super().__init__()
        self.embed_size = embed_size
        self.heads = heads
        self.head_dim = embed_size // heads
        
        self.queries = nn.Linear(embed_size, embed_size)
        self.keys = nn.Linear(embed_size, embed_size)
        self.values = nn.Linear(embed_size, embed_size)
        self.fc_out = nn.Linear(embed_size, embed_size)
    
    def forward(self, x, mask=None):
        N, seq_length, _ = x.shape
        q = self.queries(x).view(N, seq_length, self.heads, self.head_dim)
        k = self.keys(x).view(N, seq_length, self.heads, self.head_dim)
        v = self.values(x).view(N, seq_length, self.heads, self.head_dim)
        
        scores = torch.einsum("nqhd,nkhd->nhqk", q, k) / math.sqrt(self.head_dim)
        if mask is not None:
            scores = scores.masked_fill(mask == 0, float("-inf"))
        attention = torch.softmax(scores, dim=-1)
        out = torch.einsum("nhql,nlhd->nqhd", attention, v)
        out = out.reshape(N, seq_length, self.embed_size)
        return self.fc_out(out)
''',
    },
}

# Diagram data for specific topics
DIAGRAM_DATA: Dict[str, Dict[str, Any]] = {
    "neural network": {
        "flowchart": {
            "title": "Neural Network Forward Pass",
            "steps": ["Input Layer", "Hidden Layer 1\\n(ReLU)", "Hidden Layer 2\\n(ReLU)", "Output Layer\\n(Softmax)"],
        },
        "architecture": {
            "components": ["Input (784)", "Dense (256)", "Dropout", "Dense (128)", "Output (10)"],
            "connections": [(0, 1), (1, 2), (2, 3), (3, 4)],
        },
    },
    "machine learning": {
        "flowchart": {
            "title": "ML Pipeline",
            "steps": ["Data Collection", "Preprocessing", "Feature Engineering", "Model Training", "Evaluation", "Deployment"],
        },
    },
    "deep learning": {
        "flowchart": {
            "title": "Deep Learning Training Loop",
            "steps": ["Forward Pass", "Compute Loss", "Backpropagation", "Update Weights", "Repeat"],
        },
        "architecture": {
            "components": ["Input", "Conv2D", "ReLU", "MaxPool", "Conv2D", "ReLU", "MaxPool", "Flatten", "Dense", "Output"],
            "connections": [(0, 1), (1, 2), (2, 3), (3, 4), (4, 5), (5, 6), (6, 7), (7, 8), (8, 9)],
        },
    },
    "transformer": {
        "architecture": {
            "components": ["Input Embedding", "Positional Encoding", "Multi-Head Attention", "Add & Norm", "Feed Forward", "Add & Norm", "Output"],
            "connections": [(0, 1), (1, 2), (2, 3), (3, 4), (4, 5), (5, 6)],
        },
    },
    "gradient descent": {
        "flowchart": {
            "title": "Gradient Descent Algorithm",
            "steps": ["Initialize Weights", "Compute Prediction", "Calculate Loss", "Compute Gradients", "Update Weights", "Converged?"],
        },
    },
    "generative adversarial network": {
        "architecture": {
            "components": ["Noise Vector (z)", "Generator (G)", "Fake Data", "Discriminator (D)", "Real/Fake Classification", "Real Data"],
            "connections": [(0, 1), (1, 2), (2, 3), (5, 3), (3, 4)],
        },
    },
}


# ──────────────────────────────────────────────────────────────────────────────
# Utility Functions
# ──────────────────────────────────────────────────────────────────────────────

def _detect_language(text: str) -> str:
    """Auto-detect language from text content."""
    if not text:
        return "en"
    for ch in text:
        if ch in HEBREW_ALPHABET:
            return "he"
    return "en"


def _get_font_path() -> Optional[str]:
    """Find an available font file path."""
    for fp in FONT_PATHS:
        if os.path.isfile(fp):
            return fp
    return None


def _ensure_dir(path: Path) -> Path:
    """Ensure directory exists, create if not."""
    path.mkdir(parents=True, exist_ok=True)
    return path


def _hash_topic(topic: str) -> str:
    """Generate a short hash from a topic string."""
    return hashlib.md5(topic.encode("utf-8")).hexdigest()[:8]


def _get_output_path(topic: str, modality: str, ext: str) -> str:
    """Generate a timestamped output file path."""
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    h = _hash_topic(topic)
    out_dir = _ensure_dir(DEFAULT_OUTPUT_DIR / modality)
    return str(out_dir / f"{modality}_{h}_{ts}.{ext}")


def _find_knowledge(topic: str) -> Optional[str]:
    """Find a knowledge base entry matching the topic."""
    topic_lower = topic.lower()
    for key in KNOWLEDGE_BASE:
        if key in topic_lower:
            lang = _detect_language(topic)
            return KNOWLEDGE_BASE[key].get(lang, KNOWLEDGE_BASE[key].get("en", ""))
    return None


def _find_diagram_data(topic: str, diagram_type: str) -> Optional[Dict[str, Any]]:
    """Find diagram data for a topic."""
    topic_lower = topic.lower()
    for key in DIAGRAM_DATA:
        if key in topic_lower and diagram_type in DIAGRAM_DATA[key]:
            return DIAGRAM_DATA[key][diagram_type]
    return None


def _find_code_example(topic: str, language: str) -> Optional[str]:
    """Find a code example for a topic."""
    topic_lower = topic.lower()
    for key in CODE_EXAMPLES:
        if key in topic_lower and language in CODE_EXAMPLES[key]:
            return CODE_EXAMPLES[key][language]
    return None


# ──────────────────────────────────────────────────────────────────────────────
# MultimodalOutput Dataclass
# ──────────────────────────────────────────────────────────────────────────────

@dataclass
class MultimodalOutput:
    """Unified output containing multiple modalities.

    Attributes:
        text: Textual explanation of the topic.
        image_path: Path to generated diagram/image file.
        audio_path: Path to TTS audio file.
        document_path: Path to generated PDF/Word document.
        code_blocks: List of code snippets as strings.
        diagram_svg: SVG diagram markup as a string.
        metadata: Additional metadata dict (timestamps, language, etc.).
    """

    text: str = ""
    image_path: Optional[str] = None
    audio_path: Optional[str] = None
    document_path: Optional[str] = None
    code_blocks: List[str] = field(default_factory=list)
    diagram_svg: Optional[str] = None
    metadata: Dict[str, Any] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to plain dictionary."""
        return asdict(self)

    def to_json(self, indent: int = 2) -> str:
        """Serialize to JSON string."""
        return json.dumps(self.to_dict(), indent=indent, ensure_ascii=False, default=str)

    def has_modality(self, modality: str) -> bool:
        """Check whether a specific modality is present.

        Supported modalities:
            - "text"    → self.text is non-empty
            - "image"   → self.image_path is set
            - "audio"   → self.audio_path is set
            - "document"→ self.document_path is set
            - "code"    → self.code_blocks is non-empty
            - "diagram" → self.diagram_svg is set
        """
        modality = modality.lower().strip()
        mapping = {
            "text": bool(self.text),
            "image": self.image_path is not None,
            "audio": self.audio_path is not None,
            "document": self.document_path is not None,
            "code": bool(self.code_blocks),
            "diagram": self.diagram_svg is not None,
        }
        return mapping.get(modality, False)

    def summary(self) -> str:
        """Return a human-readable summary of what was generated."""
        parts = ["Multimodal Output Summary:"]
        parts.append(f"  Text: {'Yes (' + str(len(self.text)) + ' chars)' if self.text else 'No'}")
        parts.append(f"  Image: {'Yes (' + self.image_path + ')' if self.image_path else 'No'}")
        parts.append(f"  Audio: {'Yes (' + self.audio_path + ')' if self.audio_path else 'No'}")
        parts.append(f"  Document: {'Yes (' + self.document_path + ')' if self.document_path else 'No'}")
        parts.append(f"  Code blocks: {len(self.code_blocks)}")
        parts.append(f"  Diagram SVG: {'Yes' if self.diagram_svg else 'No'}")
        parts.append(f"  Metadata keys: {list(self.metadata.keys())}")
        return "\n".join(parts)

    def __repr__(self) -> str:
        mods = []
        if self.text:
            mods.append("text")
        if self.image_path:
            mods.append("image")
        if self.audio_path:
            mods.append("audio")
        if self.document_path:
            mods.append("document")
        if self.code_blocks:
            mods.append("code")
        if self.diagram_svg:
            mods.append("diagram")
        return f"MultimodalOutput(modalities={mods}, topic={self.metadata.get('topic', 'N/A')!r})"


# ──────────────────────────────────────────────────────────────────────────────
# Diagram Generation — Local (matplotlib), no external APIs
# ──────────────────────────────────────────────────────────────────────────────

def _create_flowchart(title: str, steps: List[str]) -> str:
    """Create a flowchart diagram using matplotlib.

    Args:
        title: Chart title.
        steps: List of step labels.

    Returns:
        File path to the saved PNG image.
    """
    if not _HAS_MATPLOTLIB:
        raise RuntimeError("matplotlib is required for flowchart generation.")

    if not steps:
        steps = ["Start", "Process", "End"]

    fig, ax = plt.subplots(figsize=(max(8, len(steps) * 2.5), 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, len(steps) * 1.5 + 1)
    ax.axis("off")
    ax.set_title(title, fontsize=16, fontweight="bold", pad=20)

    box_w, box_h = 1.8, 0.8
    y_positions = [len(steps) * 1.5 - i * 1.5 for i in range(len(steps))]

    colors_seq = plt.cm.Blues(np.linspace(0.4, 0.9, len(steps))) if _HAS_NUMPY else ["#5B9BD5"] * len(steps)

    for i, (step, y) in enumerate(zip(steps, y_positions)):
        x = 5 - box_w / 2
        color = colors_seq[i] if _HAS_NUMPY else "#5B9BD5"
        box = FancyBboxPatch(
            (x, y), box_w, box_h,
            boxstyle="round,pad=0.15",
            facecolor=color,
            edgecolor="#2B547E",
            linewidth=1.5,
        )
        ax.add_patch(box)
        ax.text(5, y + box_h / 2, step, ha="center", va="center",
                fontsize=10, fontweight="bold", color="white", wrap=True)

        # Draw arrow to next step
        if i < len(steps) - 1:
            ax.annotate("", xy=(5, y - 0.2), xytext=(5, y - 0.7),
                        arrowprops=dict(arrowstyle="->", color="#2B547E", lw=2))

    # Add decision diamond for last step if it looks like a question
    if steps and "?" in steps[-1]:
        y = y_positions[-1]
        # Redraw last as diamond
        diamond = plt.Polygon(
            [[5, y + box_h + 0.3], [5 + box_w / 2 + 0.2, y + box_h / 2],
             [5, y - 0.3], [5 - box_w / 2 - 0.2, y + box_h / 2]],
            facecolor="#E74C3C", edgecolor="#922B21", linewidth=1.5
        )
        ax.add_patch(diamond)
        ax.text(5, y + box_h / 2, steps[-1], ha="center", va="center",
                fontsize=9, fontweight="bold", color="white")

    out_path = _get_output_path(title, "diagrams", "png")
    plt.tight_layout()
    fig.savefig(out_path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out_path


def _create_mindmap(title: str, nodes: Dict[str, List[str]]) -> str:
    """Create a mind map using matplotlib/networkx.

    Args:
        title: Mind map title.
        nodes: Dict with center node as key and list of branches as values.
               e.g. {"AI": ["ML", "DL", "NLP", "CV"]}

    Returns:
        File path to the saved PNG image.
    """
    if not _HAS_MATPLOTLIB:
        raise RuntimeError("matplotlib is required for mind map generation.")

    if _HAS_NETWORKX:
        G = nx.Graph()
        center = list(nodes.keys())[0] if nodes else "Topic"
        branches = nodes.get(center, ["A", "B", "C"])

        G.add_node(center)
        for branch in branches:
            G.add_edge(center, branch)
            if isinstance(nodes.get(branch), list):
                for leaf in nodes[branch]:
                    G.add_edge(branch, leaf)

        fig, ax = plt.subplots(figsize=(12, 10))
        pos = nx.spring_layout(G, k=2, iterations=50, seed=42)

        # Node colors
        node_colors = []
        for n in G.nodes():
            if n == center:
                node_colors.append("#E74C3C")
            elif n in branches:
                node_colors.append("#3498DB")
            else:
                node_colors.append("#2ECC71")

        nx.draw_networkx_nodes(G, pos, node_color=node_colors, node_size=3000, ax=ax,
                               edgecolors="white", linewidths=2)
        nx.draw_networkx_edges(G, pos, edge_color="#95A5A6", width=2, alpha=0.7, ax=ax)
        nx.draw_networkx_labels(G, pos, font_size=10, font_weight="bold", ax=ax)

        ax.set_title(title, fontsize=16, fontweight="bold", pad=20)
        ax.axis("off")

    else:
        # Fallback without networkx — simple radial layout
        fig, ax = plt.subplots(figsize=(12, 10))
        ax.set_xlim(-6, 6)
        ax.set_ylim(-6, 6)
        ax.set_aspect("equal")
        ax.axis("off")
        ax.set_title(title, fontsize=16, fontweight="bold", pad=20)

        center = list(nodes.keys())[0] if nodes else "Topic"
        branches = nodes.get(center, ["A", "B", "C"])

        # Draw center
        center_circle = plt.Circle((0, 0), 1.2, facecolor="#E74C3C", edgecolor="#922B21", linewidth=2)
        ax.add_patch(center_circle)
        ax.text(0, 0, center, ha="center", va="center", fontsize=12, fontweight="bold", color="white")

        # Draw branches radially
        n = len(branches)
        for i, branch in enumerate(branches):
            angle = 2 * 3.14159 * i / n
            x, y = 3.5 * np.cos(angle), 3.5 * np.sin(angle)
            if not _HAS_NUMPY:
                import math
                x, y = 3.5 * math.cos(angle), 3.5 * math.sin(angle)

            # Connection line
            ax.plot([1.0 * np.cos(angle) if _HAS_NUMPY else 1.0 * math.cos(angle), x],
                    [1.0 * np.sin(angle) if _HAS_NUMPY else 1.0 * math.sin(angle), y],
                    color="#3498DB", linewidth=2, alpha=0.6)

            branch_circle = plt.Circle((x, y), 0.9, facecolor="#3498DB", edgecolor="#21618C", linewidth=1.5)
            ax.add_patch(branch_circle)
            ax.text(x, y, branch, ha="center", va="center", fontsize=9, fontweight="bold", color="white")

    out_path = _get_output_path(title, "diagrams", "png")
    plt.tight_layout()
    fig.savefig(out_path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out_path


def _create_architecture_diagram(components: List[str], connections: List[Tuple[int, int]]) -> str:
    """Create an architecture diagram with connected components.

    Args:
        components: List of component labels.
        connections: List of (from_index, to_index) tuples.

    Returns:
        File path to the saved PNG image.
    """
    if not _HAS_MATPLOTLIB:
        raise RuntimeError("matplotlib is required for architecture diagram generation.")

    if not components:
        components = ["Input", "Process", "Output"]
        connections = [(0, 1), (1, 2)]

    n = len(components)
    fig_w = max(12, n * 2.5)
    fig, ax = plt.subplots(figsize=(fig_w, 5))
    ax.set_xlim(-1, n * 2 + 1)
    ax.set_ylim(-1, 4)
    ax.axis("off")
    ax.set_title("Architecture Diagram", fontsize=16, fontweight="bold", pad=20)

    box_w, box_h = 1.6, 1.0
    colors = plt.cm.Set3(np.linspace(0, 1, n)) if _HAS_NUMPY else ["#98D8C8"] * n

    positions = []
    for i, comp in enumerate(components):
        x = i * 2 + 0.5
        y = 1.5
        positions.append((x + box_w / 2, y + box_h / 2))

        color = colors[i] if _HAS_NUMPY else "#98D8C8"
        box = FancyBboxPatch(
            (x, y), box_w, box_h,
            boxstyle="round,pad=0.1",
            facecolor=color,
            edgecolor="#555555",
            linewidth=1.5,
        )
        ax.add_patch(box)
        ax.text(x + box_w / 2, y + box_h / 2, comp, ha="center", va="center",
                fontsize=9, fontweight="bold", wrap=True)

    # Draw connections
    for fr, to in connections:
        if 0 <= fr < n and 0 <= to < n:
            x1, y1 = positions[fr]
            x2, y2 = positions[to]
            # Shift to edges of boxes
            x1 += box_w / 2
            x2 -= box_w / 2
            ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                        arrowprops=dict(arrowstyle="->", color="#555", lw=2,
                                       connectionstyle="arc3,rad=0.1"))

    out_path = _get_output_path("architecture", "diagrams", "png")
    plt.tight_layout()
    fig.savefig(out_path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out_path


def _create_comparison_table(items: List[str], criteria: List[str], values: List[List[str]] = None) -> str:
    """Create a comparison table/chart.

    Args:
        items: List of item names (columns).
        criteria: List of criteria names (rows).
        values: Optional 2D list of cell values.

    Returns:
        File path to the saved PNG image.
    """
    if not _HAS_MATPLOTLIB:
        raise RuntimeError("matplotlib is required for comparison table generation.")

    if not items:
        items = ["A", "B", "C"]
    if not criteria:
        criteria = ["Speed", "Accuracy", "Complexity"]
    if values is None:
        values = [["High" if (i + j) % 2 == 0 else "Low" for j in range(len(items))]
                  for i in range(len(criteria))]

    fig, ax = plt.subplots(figsize=(max(8, len(items) * 2.5), max(4, len(criteria) * 0.8 + 2)))
    ax.axis("tight")
    ax.axis("off")
    ax.set_title("Comparison Table", fontsize=14, fontweight="bold", pad=20)

    # Build table data
    table_data = [[c] + row for c, row in zip(criteria, values)]
    col_labels = ["Criterion"] + items

    table = ax.table(
        cellText=table_data,
        colLabels=col_labels,
        cellLoc="center",
        loc="center",
    )
    table.auto_set_font_size(False)
    table.set_fontsize(10)
    table.scale(1, 2)

    # Style header
    for j in range(len(col_labels)):
        table[0, j].set_facecolor("#2C3E50")
        table[0, j].set_text_props(color="white", fontweight="bold")

    # Alternate row colors
    for i in range(1, len(table_data) + 1):
        for j in range(len(col_labels)):
            color = "#ECF0F1" if i % 2 == 0 else "#FFFFFF"
            table[i, j].set_facecolor(color)
            table[i, j].set_edgecolor("#BDC3C7")

    out_path = _get_output_path("comparison", "diagrams", "png")
    plt.tight_layout()
    fig.savefig(out_path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out_path


def _create_timeline(events: List[Tuple[str, str]]) -> str:
    """Create a timeline diagram.

    Args:
        events: List of (date/label, description) tuples.

    Returns:
        File path to the saved PNG image.
    """
    if not _HAS_MATPLOTLIB:
        raise RuntimeError("matplotlib is required for timeline generation.")

    if not events:
        events = [("2020", "Start"), ("2021", "Growth"), ("2022", "Maturity"), ("2023", "Expansion")]

    fig, ax = plt.subplots(figsize=(max(10, len(events) * 3), 5))
    ax.set_xlim(-0.5, len(events) - 0.5)
    ax.set_ylim(-2, 2)
    ax.axis("off")
    ax.set_title("Timeline", fontsize=16, fontweight="bold", pad=20)

    # Draw horizontal line
    ax.hlines(0, -0.3, len(events) - 0.7, colors="#3498DB", linewidth=3)

    colors = plt.cm.Set2(np.linspace(0, 1, len(events))) if _HAS_NUMPY else ["#E74C3C"] * len(events)

    for i, (date, desc) in enumerate(events):
        color = colors[i] if _HAS_NUMPY else "#E74C3C"

        # Event dot
        ax.plot(i, 0, "o", markersize=18, color=color, zorder=5, markeredgecolor="white", markeredgewidth=2)

        # Alternating above/below
        y_offset = 1.2 if i % 2 == 0 else -1.2
        va = "bottom" if i % 2 == 0 else "top"

        # Connector line
        ax.vlines(i, 0, y_offset * 0.7, colors="#7F8C8D", linewidth=1.5, linestyles="--")

        # Date label
        ax.text(i, y_offset * 0.85, str(date), ha="center", va=va,
                fontsize=11, fontweight="bold", color=color)
        # Description
        ax.text(i, y_offset * 0.6, str(desc), ha="center", va=va,
                fontsize=9, color="#2C3E50", wrap=True)

    out_path = _get_output_path("timeline", "diagrams", "png")
    plt.tight_layout()
    fig.savefig(out_path, dpi=150, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return out_path


# ──────────────────────────────────────────────────────────────────────────────
# Document Generation
# ──────────────────────────────────────────────────────────────────────────────

def _generate_pdf(title: str, content: str, output_path: str) -> str:
    """Generate a PDF document using ReportLab."""
    if not _HAS_REPORTLAB:
        raise RuntimeError("reportlab is required for PDF generation.")

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=18,
    )

    styles = getSampleStyleSheet()
    story: List[Any] = []

    # Title
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Heading1"],
        fontSize=20,
        textColor=colors.HexColor("#2C3E50"),
        spaceAfter=20,
        alignment=TA_CENTER,
    )
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 0.2 * inch))

    # Date
    date_style = ParagraphStyle(
        "DateStyle",
        parent=styles["Normal"],
        fontSize=10,
        textColor=colors.grey,
        alignment=TA_CENTER,
    )
    story.append(Paragraph(
        f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}",
        date_style
    ))
    story.append(Spacer(1, 0.3 * inch))

    # Content
    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["Normal"],
        fontSize=11,
        leading=16,
        spaceAfter=12,
    )

    for paragraph in content.split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        if paragraph.startswith("# "):
            story.append(Paragraph(paragraph[2:], styles["Heading1"]))
        elif paragraph.startswith("## "):
            story.append(Paragraph(paragraph[3:], styles["Heading2"]))
        elif paragraph.startswith("### "):
            story.append(Paragraph(paragraph[4:], styles["Heading3"]))
        elif paragraph.startswith("- "):
            for line in paragraph.split("\n"):
                if line.strip().startswith("- "):
                    story.append(Paragraph(f"  \u2022 {line.strip()[2:]}", body_style))
                else:
                    story.append(Paragraph(line.strip(), body_style))
        elif paragraph.startswith("```"):
            code_style = ParagraphStyle(
                "CodeStyle",
                parent=styles["Code"],
                fontSize=9,
                fontName="Courier",
                backColor=colors.HexColor("#F4F4F4"),
                leftIndent=20,
                rightIndent=20,
                spaceAfter=12,
            )
            code_lines = paragraph.split("\n")[1:-1]
            code_text = "<br/>".join(code_lines)
            story.append(Paragraph(f"<pre>{code_text}</pre>", code_style))
        else:
            story.append(Paragraph(paragraph.replace("\n", "<br/>"), body_style))

        story.append(Spacer(1, 0.1 * inch))

    doc.build(story)
    return output_path


def _generate_docx(title: str, content: str, output_path: str) -> str:
    """Generate a Word document using python-docx."""
    if not _HAS_DOCX:
        raise RuntimeError("python-docx is required for DOCX generation.")

    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(
        f"Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # Spacer

    # Content
    for paragraph in content.split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        if paragraph.startswith("# "):
            doc.add_heading(paragraph[2:], level=1)
        elif paragraph.startswith("## "):
            doc.add_heading(paragraph[3:], level=2)
        elif paragraph.startswith("### "):
            doc.add_heading(paragraph[4:], level=3)
        elif paragraph.startswith("- "):
            for line in paragraph.split("\n"):
                if line.strip().startswith("- "):
                    doc.add_paragraph(line.strip()[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line.strip())
        elif paragraph.startswith("```"):
            p = doc.add_paragraph()
            run = p.add_run("\n".join(paragraph.split("\n")[1:-1]))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(12)
        else:
            doc.add_paragraph(paragraph)

    doc.save(output_path)
    return output_path


def _generate_html(title: str, content: str, output_path: str) -> str:
    """Generate an HTML document."""
    html_template = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px;
            line-height: 1.6;
            color: #333;
            background: #f9f9f9;
        }}
        h1 {{ color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #34495e; margin-top: 30px; }}
        h3 {{ color: #555; }}
        code {{
            background: #f4f4f4;
            padding: 2px 8px;
            border-radius: 4px;
            font-family: 'Courier New', monospace;
        }}
        pre {{
            background: #2c3e50;
            color: #ecf0f1;
            padding: 20px;
            border-radius: 8px;
            overflow-x: auto;
        }}
        blockquote {{
            border-left: 4px solid #3498db;
            margin: 0;
            padding-left: 20px;
            color: #666;
        }}
        ul {{ padding-left: 20px; }}
        .meta {{
            color: #999;
            font-size: 0.9em;
            text-align: center;
            margin-bottom: 30px;
        }}
        .container {{ background: white; padding: 40px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
    </style>
</head>
<body>
    <div class="container">
        <h1>{title}</h1>
        <p class="meta">Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
        <hr>
        {_markdown_to_html(content)}
    </div>
</body>
</html>"""

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(html_template)
    return output_path


def _generate_md(title: str, content: str, output_path: str) -> str:
    """Generate a Markdown document."""
    md_content = f"""# {title}

> Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}

---

{content}
"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(md_content)
    return output_path


def _markdown_to_html(content: str) -> str:
    """Simple Markdown to HTML converter for document generation."""
    lines = content.split("\n")
    result = []
    in_code = False
    code_lines = []

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                _newline_join = "\n".join(code_lines)
                result.append(f"<pre><code>{_newline_join}</code></pre>")
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if stripped.startswith("# "):
            result.append(f"<h1>{stripped[2:]}</h1>")
        elif stripped.startswith("## "):
            result.append(f"<h2>{stripped[3:]}</h2>")
        elif stripped.startswith("### "):
            result.append(f"<h3>{stripped[4:]}</h3>")
        elif stripped.startswith("- "):
            result.append(f"<ul><li>{stripped[2:]}</li></ul>")
        elif stripped.startswith("> "):
            result.append(f"<blockquote>{stripped[2:]}</blockquote>")
        elif stripped:
            result.append(f"<p>{stripped}</p>")
        else:
            result.append("<br/>")

    return "\n".join(result)


# ──────────────────────────────────────────────────────────────────────────────
# SVG Diagram Generation
# ──────────────────────────────────────────────────────────────────────────────

def _create_svg_flowchart(title: str, steps: List[str]) -> str:
    """Create an SVG flowchart string."""
    if not steps:
        steps = ["Start", "Process", "End"]

    box_h = 50
    gap = 40
    total_h = len(steps) * (box_h + gap) + 60
    box_w = 200
    x_center = 250

    svg_parts = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="500" height="{total_h}">',
        f'<rect width="500" height="{total_h}" fill="#fafafa"/>',
        f'<text x="250" y="35" text-anchor="middle" font-size="18" font-weight="bold" fill="#2C3E50">{title}</text>',
    ]

    colors = ["#3498DB", "#2ECC71", "#E74C3C", "#9B59B6", "#F39C12", "#1ABC9C"]

    for i, step in enumerate(steps):
        y = 60 + i * (box_h + gap)
        color = colors[i % len(colors)]

        # Box
        svg_parts.append(
            f'<rect x="{x_center - box_w // 2}" y="{y}" width="{box_w}" height="{box_h}" '
            f'rx="10" fill="{color}" stroke="#2C3E50" stroke-width="2"/>'
        )
        # Text
        svg_parts.append(
            f'<text x="{x_center}" y="{y + box_h // 2 + 5}" text-anchor="middle" '
            f'font-size="14" font-weight="bold" fill="white">{step}</text>'
        )

        # Arrow
        if i < len(steps) - 1:
            arrow_y = y + box_h
            svg_parts.append(
                f'<line x1="{x_center}" y1="{arrow_y}" x2="{x_center}" y2="{arrow_y + gap}" '
                f'stroke="#2C3E50" stroke-width="2" marker-end="url(#arrowhead)"/>'
            )

    # Arrow marker definition
    svg_parts.insert(1,
        '<defs><marker id="arrowhead" markerWidth="10" markerHeight="7" '
        'refX="10" refY="3.5" orient="auto"><polygon points="0 0, 10 3.5, 0 7" fill="#2C3E50"/></marker></defs>'
    )

    svg_parts.append("</svg>")
    return "\n".join(svg_parts)


# ──────────────────────────────────────────────────────────────────────────────
# Mock Engines — used when real engines aren't available
# ──────────────────────────────────────────────────────────────────────────────

class MockTTSEngine:
    """Mock TTS engine that generates placeholder audio."""

    def synthesize(self, text: str, output_path: str, language: str = "he") -> str:
        """Create a placeholder text file instead of audio."""
        placeholder_path = output_path.rsplit(".", 1)[0] + "_placeholder.txt"
        with open(placeholder_path, "w", encoding="utf-8") as f:
            f.write(f"# Mock TTS Placeholder\nLanguage: {language}\nText: {text[:200]}...\n")
        return placeholder_path


class MockDocEngine:
    """Mock document engine that creates placeholder documents."""

    def create_document(self, title: str, content: str, output_path: str) -> str:
        """Create a placeholder text file."""
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(f"# {title}\n\n{content}\n")
        return output_path


class MockDrawEngine:
    """Mock drawing engine that creates placeholder images."""

    def draw(self, description: str, output_path: str) -> str:
        """Create a placeholder blank image if PIL is available, else a text file."""
        if _HAS_PIL:
            img = Image.new("RGB", (800, 600), color=(240, 240, 240))
            draw = ImageDraw.Draw(img)
            draw.text((400, 300), f"[Mock Image] - {description[:50]}",
                     fill=(100, 100, 100), anchor="mm")
            img.save(output_path)
            return output_path
        else:
            with open(output_path, "w") as f:
                f.write(f"[Mock Image Placeholder] {description}\n")
            return output_path


# ──────────────────────────────────────────────────────────────────────────────
# MultimodalEngine — Main Class
# ──────────────────────────────────────────────────────────────────────────────

class MultimodalEngine:
    """Generates multimodal outputs from a single request.

    Combines: text + image + voice + document + code + diagram.
    All processing is 100% local with optional external integrations.

    Args:
        voice_engine: Optional TTS engine (gTTS, pyttsx3, or custom).
        doc_engine: Optional document generation engine.
        draw_engine: Optional drawing/image generation engine.
        output_dir: Base directory for all generated files.
    """

    def __init__(
        self,
        voice_engine: Any = None,
        doc_engine: Any = None,
        draw_engine: Any = None,
        output_dir: Optional[str] = None,
    ):
        self.voice_engine = voice_engine or (MockTTSEngine() if not (_HAS_GTTS or _HAS_PYTTSX3) else None)
        self.doc_engine = doc_engine or (MockDocEngine() if not (_HAS_REPORTLAB or _HAS_DOCX) else None)
        self.draw_engine = draw_engine or MockDrawEngine()
        self.output_dir = Path(output_dir) if output_dir else DEFAULT_OUTPUT_DIR
        self._capabilities: Optional[List[str]] = None

        # Ensure output directories exist
        for sub in ["diagrams", "audio", "documents", "images", "code"]:
            _ensure_dir(self.output_dir / sub)

    # ── Main explain method ─────────────────────────────────────────────────

    def explain(
        self,
        topic: str,
        modalities: Optional[List[str]] = None,
        language: str = "he",
    ) -> MultimodalOutput:
        """Explain a topic using all requested modalities.

        Args:
            topic: The topic to explain.
            modalities: List of modalities to generate. Defaults to ["text", "diagram"].
            language: Primary language ("he" for Hebrew, "en" for English).

        Returns:
            MultimodalOutput with all generated content.
        """
        if modalities is None:
            modalities = ["text", "diagram"]
        modalities = [m.lower().strip() for m in modalities]
        topic_lower = topic.lower()

        output = MultimodalOutput(
            metadata={
                "topic": topic,
                "modalities": modalities,
                "language": language,
                "timestamp": datetime.datetime.now().isoformat(),
                "engine": "MultimodalEngine",
            }
        )

        # 1. Text explanation
        if "text" in modalities:
            output.text = self.generate_text_explanation(topic, language)

        # 2. Diagram / Image
        if "diagram" in modalities or "image" in modalities:
            try:
                diagram_type = "flowchart"
                if any(t in topic_lower for t in ["architecture", "system", "structure"]):
                    diagram_type = "architecture"
                elif any(t in topic_lower for t in ["compare", "comparison", "vs", "versus"]):
                    diagram_type = "comparison"
                elif any(t in topic_lower for t in ["history", "timeline", "evolution"]):
                    diagram_type = "timeline"
                elif any(t in topic_lower for t in ["mind", "concept", "map"]):
                    diagram_type = "mindmap"

                if "diagram" in modalities:
                    output.diagram_svg = self.generate_diagram(topic, diagram_type)
                if "image" in modalities:
                    output.image_path = self.generate_diagram(topic, diagram_type)
            except Exception as e:
                output.metadata["diagram_error"] = str(e)

        # 3. Code example
        if "code" in modalities:
            code = self.generate_code_example(topic, language="python")
            if code:
                output.code_blocks.append(code)

        # 4. Audio
        if "audio" in modalities:
            text_for_audio = output.text or self.generate_text_explanation(topic, language)
            try:
                output.audio_path = self.generate_audio_explanation(text_for_audio, language)
            except Exception as e:
                output.metadata["audio_error"] = str(e)

        # 5. Document
        if "document" in modalities:
            doc_content = output.text or self.generate_text_explanation(topic, language)
            doc_content += "\n\n"
            if output.code_blocks:
                doc_content += "## Code Example\n\n```python\n" + output.code_blocks[0] + "\n```\n"
            try:
                output.document_path = self.generate_document(topic, doc_content, doc_type="pdf")
            except Exception as e:
                output.metadata["document_error"] = str(e)

        output.metadata["completed_at"] = datetime.datetime.now().isoformat()
        return output

    # ── Individual generation methods ────────────────────────────────────────

    def generate_text_explanation(self, topic: str, language: str = "he") -> str:
        """Generate a textual explanation for a topic.

        Uses the built-in knowledge base for known topics, or generates
        a template-based explanation for unknown ones.

        Args:
            topic: The topic to explain.
            language: "he" for Hebrew, "en" for English.

        Returns:
            Textual explanation string.
        """
        # Try knowledge base
        kb_entry = _find_knowledge(topic)
        if kb_entry:
            return kb_entry

        # Auto-detect language from topic text
        detected = _detect_language(topic)
        if detected != language and language == "he":
            language = detected

        # Generate a template-based explanation
        if language == "he":
            return (
                f"**{topic}** הוא נושא מרכזי בתחום הבינה המלאכותית והמדע. "
                f"המושג {topic} מתייחס למערכת או שיטה המאפשרת למחשבים לבצע משימות "
                f"הדורשות בינה אנושית כמו למידה, חיזוי, או קבלת החלטות. "
                f"הטכנולוגיה מבוססת על אלגוריתמים מתקדמים המעבדים כמויות גדולות של נתונים "
                f"ומזהים בהם דפוסים על מנת לשפר ביצועים לאורך זמן. "
                f"היישומים של {topic} נפוצים בתעשיות רבות כולל רפואה, פיננסים, "
                f"תחבורה, ובידור. הבנת {topic} חיונית לכל מי שמתעניין בעתיד הטכנולוגיה."
            )
        else:
            return (
                f"**{topic}** is a key concept in artificial intelligence and computer science. "
                f"It refers to a system or methodology that enables computers to perform tasks "
                f"that typically require human intelligence, such as learning, reasoning, and "
                f"decision-making. The technology relies on advanced algorithms that process "
                f"large amounts of data and identify patterns to improve performance over time. "
                f"Applications of {topic} span numerous industries including healthcare, "
                f"finance, transportation, and entertainment. Understanding {topic} is "
                f"essential for anyone interested in the future of technology."
            )

    def generate_diagram(self, topic: str, diagram_type: str = "flowchart") -> str:
        """Create a visual diagram for a topic.

        Args:
            topic: The topic to diagram.
            diagram_type: One of "flowchart", "mindmap", "timeline",
                          "architecture", "comparison".

        Returns:
            File path (or SVG string) of the generated diagram.
        """
        if diagram_type not in DIAGRAM_TYPES:
            diagram_type = "flowchart"

        topic_lower = topic.lower()

        # Try to find pre-built diagram data
        diagram_data = _find_diagram_data(topic, diagram_type)

        if diagram_type == "flowchart":
            if diagram_data and "steps" in diagram_data:
                steps = diagram_data["steps"]
                title = diagram_data.get("title", topic)
            else:
                steps = ["Input", "Processing", "Output"]
                if "neural" in topic_lower or "network" in topic_lower:
                    steps = ["Input Layer", "Hidden Layer (ReLU)", "Output Layer (Softmax)"]
                elif "gradient" in topic_lower:
                    steps = ["Initialize", "Forward Pass", "Compute Loss", "Backpropagate", "Update Weights"]
                elif "train" in topic_lower:
                    steps = ["Load Data", "Preprocess", "Build Model", "Train", "Evaluate"]
                title = topic
            return _create_flowchart(title, steps)

        elif diagram_type == "mindmap":
            if diagram_data:
                nodes = {topic: diagram_data.get("branches", ["Concept A", "Concept B", "Concept C"])}
            else:
                nodes = {topic: ["Overview", "Applications", "Methods", "Challenges", "Future"]}
            return _create_mindmap(topic, nodes)

        elif diagram_type == "architecture":
            if diagram_data and "components" in diagram_data:
                components = diagram_data["components"]
                connections = diagram_data.get("connections", [])
            else:
                components = ["Input", "Embedding", "Encoder", "Decoder", "Output"]
                connections = [(0, 1), (1, 2), (2, 3), (3, 4)]
                if "neural" in topic_lower:
                    components = ["Input (784)", "Dense (256)", "ReLU", "Dense (128)", "Output (10)"]
                    connections = [(0, 1), (1, 2), (2, 3), (3, 4)]
            return _create_architecture_diagram(components, connections)

        elif diagram_type == "comparison":
            items = ["Method A", "Method B", "Method C"]
            criteria = ["Speed", "Accuracy", "Complexity", "Scalability"]
            values = [["Fast", "Medium", "Slow"],
                      ["85%", "92%", "78%"],
                      ["Low", "High", "Medium"],
                      ["Good", "Excellent", "Fair"]]
            return _create_comparison_table(items, criteria, values)

        elif diagram_type == "timeline":
            if diagram_data and "events" in diagram_data:
                events = diagram_data["events"]
            else:
                events = [
                    ("1950", "AI Birth"),
                    ("1980", "Expert Systems"),
                    ("2012", "Deep Learning Era"),
                    ("2017", "Transformers"),
                    ("2023", "LLMs & GenAI"),
                ]
            return _create_timeline(events)

        # Fallback
        return _create_flowchart(topic, ["Start", "Process", "End"])

    def generate_audio_explanation(self, text: str, language: str = "he") -> str:
        """Convert text to speech and save as audio file.

        Args:
            text: Text to convert to speech.
            language: Language code ("he", "en", etc.).

        Returns:
            File path to the audio file.
        """
        output_path = _get_output_path(text[:30], "audio", "mp3")

        # Try pyttsx3 (offline)
        if _HAS_PYTTSX3 and self.voice_engine is None:
            try:
                engine = pyttsx3.init()
                engine.setProperty("rate", 150)
                engine.setProperty("volume", 0.9)
                # Try to set a voice for the language
                voices = engine.getProperty("voices")
                for v in voices:
                    if language in v.languages or language in v.id:
                        engine.setProperty("voice", v.id)
                        break
                engine.save_to_file(text, output_path)
                engine.runAndWait()
                return output_path
            except Exception:
                pass

        # Try gTTS (online, but lightweight)
        if _HAS_GTTS and self.voice_engine is None:
            try:
                lang_map = {"he": "iw", "en": "en", "ar": "ar", "fr": "fr", "es": "es"}
                tts_lang = lang_map.get(language, language)
                tts = gTTS(text=text[:5000], lang=tts_lang, slow=False)
                tts.save(output_path)
                return output_path
            except Exception:
                pass

        # Fallback to mock
        if self.voice_engine:
            return self.voice_engine.synthesize(text, output_path, language)

        # Last resort — write text file
        txt_path = output_path.replace(".mp3", ".txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(text)
        return txt_path

    def generate_document(
        self,
        title: str,
        content: str,
        doc_type: str = "pdf",
    ) -> str:
        """Create a formatted document.

        Args:
            title: Document title.
            content: Document body content (supports simple Markdown).
            doc_type: One of "pdf", "docx", "html", "md".

        Returns:
            File path to the generated document.
        """
        if doc_type not in DOCUMENT_TYPES:
            doc_type = "pdf"

        ext_map = {"pdf": "pdf", "docx": "docx", "html": "html", "md": "md"}
        output_path = _get_output_path(title, "documents", ext_map[doc_type])

        if doc_type == "pdf":
            if _HAS_REPORTLAB:
                return _generate_pdf(title, content, output_path)
            else:
                # Fallback to HTML then try conversion
                html_path = output_path.replace(".pdf", ".html")
                _generate_html(title, content, html_path)
                return html_path

        elif doc_type == "docx":
            if _HAS_DOCX:
                return _generate_docx(title, content, output_path)
            else:
                md_path = output_path.replace(".docx", ".md")
                return _generate_md(title, content, md_path)

        elif doc_type == "html":
            return _generate_html(title, content, output_path)

        elif doc_type == "md":
            return _generate_md(title, content, output_path)

        return output_path

    def generate_code_example(self, topic: str, language: str = "python") -> str:
        """Generate a code example for a topic.

        Args:
            topic: The topic to generate code for.
            language: Programming language (currently only "python" supported).

        Returns:
            Code string.
        """
        # Try database
        code = _find_code_example(topic, language)
        if code:
            return code

        topic_lower = topic.lower()

        # Generate template code
        if language == "python":
            if "neural" in topic_lower or "network" in topic_lower:
                return '''
import torch
import torch.nn as nn

class SimpleNN(nn.Module):
    def __init__(self, input_size, hidden_size, num_classes):
        super(SimpleNN, self).__init__()
        self.layer1 = nn.Linear(input_size, hidden_size)
        self.relu = nn.ReLU()
        self.layer2 = nn.Linear(hidden_size, num_classes)

    def forward(self, x):
        out = self.layer1(x)
        out = self.relu(out)
        out = self.layer2(out)
        return out

# Create model
model = SimpleNN(784, 256, 10)
print(model)
'''
            elif "gradient" in topic_lower:
                return '''
import numpy as np

def gradient_descent(f, df, x0, lr=0.1, epochs=100):
    """Simple gradient descent optimizer."""
    x = x0
    history = [x]
    for _ in range(epochs):
        grad = df(x)
        x = x - lr * grad
        history.append(x)
    return x, history

# Example: minimize f(x) = x^2
f = lambda x: x**2
df = lambda x: 2*x
minimum, path = gradient_descent(f, df, x0=5.0)
print(f"Minimum at x = {minimum:.4f}")
'''
            else:
                return '''
# Example code for: {topic}

def main():
    """Main function demonstrating the concept."""
    print("Hello, World!")
    data = [1, 2, 3, 4, 5]
    result = sum(data) / len(data)
    print(f"Average: {result}")
    return result

if __name__ == "__main__":
    main()
'''.replace("{topic}", topic)

        return f"// Code example for {topic} in {language}\n"

    def combine_outputs(self, outputs: Dict[str, Any]) -> MultimodalOutput:
        """Combine individual outputs into a unified MultimodalOutput.

        Args:
            outputs: Dict with keys like "text", "image_path", "audio_path",
                     "document_path", "code_blocks", "diagram_svg".

        Returns:
            MultimodalOutput instance.
        """
        return MultimodalOutput(
            text=outputs.get("text", ""),
            image_path=outputs.get("image_path"),
            audio_path=outputs.get("audio_path"),
            document_path=outputs.get("document_path"),
            code_blocks=outputs.get("code_blocks", []),
            diagram_svg=outputs.get("diagram_svg"),
            metadata=outputs.get("metadata", {}),
        )

    def present(self, output: MultimodalOutput, display_mode: str = "all") -> Dict[str, Any]:
        """Present a MultimodalOutput to the user.

        Args:
            output: The MultimodalOutput to present.
            display_mode: One of "all", "text_only", "visual_only",
                         "audio_only", "interactive".

        Returns:
            Dict with presentation results.
        """
        display_mode = display_mode.lower().strip()
        results = {"mode": display_mode, "actions": []}

        if display_mode in ("all", "text_only", "interactive"):
            if output.text:
                print("=" * 60)
                print("TEXT EXPLANATION")
                print("=" * 60)
                print(output.text)
                print()
                results["actions"].append("printed_text")

        if display_mode in ("all", "visual_only", "interactive"):
            if output.diagram_svg:
                svg_path = _get_output_path("diagram", "diagrams", "svg")
                with open(svg_path, "w", encoding="utf-8") as f:
                    f.write(output.diagram_svg)
                results["actions"].append(f"saved_diagram_svg:{svg_path}")

            if output.image_path and os.path.exists(output.image_path):
                results["actions"].append(f"displayed_image:{output.image_path}")

        if display_mode in ("all", "audio_only"):
            if output.audio_path and os.path.exists(output.audio_path):
                results["actions"].append(f"audio_ready:{output.audio_path}")

        if display_mode == "interactive":
            results["actions"].append("interactive_mode_enabled")

        if output.document_path:
            results["actions"].append(f"document:{output.document_path}")

        return results

    def get_capabilities(self) -> List[str]:
        """List available modalities based on installed dependencies.

        Returns:
            List of available modality strings.
        """
        if self._capabilities is not None:
            return self._capabilities

        caps = ["text", "code"]  # Always available

        if _HAS_MATPLOTLIB:
            caps.extend(["diagram", "image"])

        if _HAS_GTTS or _HAS_PYTTSX3:
            caps.append("audio")

        if _HAS_REPORTLAB:
            caps.append("document:pdf")
        if _HAS_DOCX:
            caps.append("document:docx")
        caps.append("document:html")
        caps.append("document:md")

        self._capabilities = caps
        return caps

    def __repr__(self) -> str:
        caps = self.get_capabilities()
        return f"MultimodalEngine(capabilities={caps})"


# ──────────────────────────────────────────────────────────────────────────────
# MockMultimodalEngine — Same interface, placeholder outputs
# ──────────────────────────────────────────────────────────────────────────────

class MockMultimodalEngine:
    """Mock multimodal engine with the same interface as MultimodalEngine.

    Creates placeholder outputs for all modalities.
    Useful for testing or when no dependencies are installed.
    """

    def __init__(
        self,
        voice_engine: Any = None,
        doc_engine: Any = None,
        draw_engine: Any = None,
        output_dir: Optional[str] = None,
    ):
        self.output_dir = Path(output_dir) if output_dir else DEFAULT_OUTPUT_DIR
        self.voice_engine = MockTTSEngine()
        self.doc_engine = MockDocEngine()
        self.draw_engine = MockDrawEngine()

    def explain(
        self,
        topic: str,
        modalities: Optional[List[str]] = None,
        language: str = "he",
    ) -> MultimodalOutput:
        """Generate mock explanation with placeholder outputs."""
        if modalities is None:
            modalities = ["text"]

        output = MultimodalOutput(
            metadata={
                "topic": topic,
                "modalities": modalities,
                "language": language,
                "mock": True,
                "timestamp": datetime.datetime.now().isoformat(),
            }
        )

        # Always provide text
        output.text = f"[MOCK] Explanation for '{topic}': This is a placeholder text explanation."
        if language == "he":
            output.text = f"[MOCK] הסבר לנושא '{topic}': זהו טקסט placeholder."

        if "diagram" in modalities:
            output.diagram_svg = _create_svg_flowchart(topic, ["Step 1", "Step 2", "Step 3"])

        if "image" in modalities:
            output.image_path = self.draw_engine.draw(topic, _get_output_path(topic, "images", "png"))

        if "code" in modalities:
            output.code_blocks.append(f"# Mock code for {topic}\nprint('Hello')\n")

        if "audio" in modalities:
            output.audio_path = self.voice_engine.synthesize(
                output.text,
                _get_output_path(topic, "audio", "mp3"),
                language,
            )

        if "document" in modalities:
            output.document_path = self.doc_engine.create_document(
                topic, output.text,
                _get_output_path(topic, "documents", "txt"),
            )

        return output

    def generate_text_explanation(self, topic: str, language: str = "he") -> str:
        """Return placeholder text."""
        if language == "he":
            return f"[MOCK] הסבר לנושא: {topic}"
        return f"[MOCK] Explanation for: {topic}"

    def generate_diagram(self, topic: str, diagram_type: str = "flowchart") -> str:
        """Return placeholder SVG."""
        return _create_svg_flowchart(topic, ["Start", "Process", "End"])

    def generate_audio_explanation(self, text: str, language: str = "he") -> str:
        """Return placeholder audio path."""
        return self.voice_engine.synthesize(text, _get_output_path(text[:20], "audio", "mp3"), language)

    def generate_document(self, title: str, content: str, doc_type: str = "pdf") -> str:
        """Return placeholder document path."""
        path = _get_output_path(title, "documents", doc_type if doc_type != "pdf" else "txt")
        return self.doc_engine.create_document(title, content, path)

    def generate_code_example(self, topic: str, language: str = "python") -> str:
        """Return placeholder code."""
        return f"# Mock code example for: {topic}\nprint('Hello, {topic}!')\n"

    def combine_outputs(self, outputs: Dict[str, Any]) -> MultimodalOutput:
        """Combine outputs into MultimodalOutput."""
        return MultimodalOutput(
            text=outputs.get("text", ""),
            image_path=outputs.get("image_path"),
            audio_path=outputs.get("audio_path"),
            document_path=outputs.get("document_path"),
            code_blocks=outputs.get("code_blocks", []),
            diagram_svg=outputs.get("diagram_svg"),
            metadata=outputs.get("metadata", {}),
        )

    def present(self, output: MultimodalOutput, display_mode: str = "all") -> Dict[str, Any]:
        """Mock present method."""
        return {"mode": display_mode, "actions": ["mock_present"]}

    def get_capabilities(self) -> List[str]:
        """Return mock capabilities."""
        return ["text", "diagram", "code"]

    def __repr__(self) -> str:
        return "MockMultimodalEngine()"


# ──────────────────────────────────────────────────────────────────────────────
# Factory Function
# ──────────────────────────────────────────────────────────────────────────────

def get_multimodal_engine(
    use_mock: bool = False,
    voice_engine: Any = None,
    doc_engine: Any = None,
    draw_engine: Any = None,
    output_dir: Optional[str] = None,
) -> MultimodalEngine:
    """Factory function to create a MultimodalEngine (or MockMultimodalEngine).

    Args:
        use_mock: Force use of MockMultimodalEngine.
        voice_engine: Optional voice engine.
        doc_engine: Optional document engine.
        draw_engine: Optional drawing engine.
        output_dir: Output directory path.

    Returns:
        MultimodalEngine or MockMultimodalEngine instance.
    """
    if use_mock:
        return MockMultimodalEngine(
            voice_engine=voice_engine,
            doc_engine=doc_engine,
            draw_engine=draw_engine,
            output_dir=output_dir,
        )

    # Check if we have enough real capabilities
    has_minimal = _HAS_MATPLOTLIB
    if not has_minimal:
        return MockMultimodalEngine(
            voice_engine=voice_engine,
            doc_engine=doc_engine,
            draw_engine=draw_engine,
            output_dir=output_dir,
        )

    return MultimodalEngine(
        voice_engine=voice_engine,
        doc_engine=doc_engine,
        draw_engine=draw_engine,
        output_dir=output_dir,
    )


# ──────────────────────────────────────────────────────────────────────────────
# Command-line interface
# ──────────────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python multimodal_output.py <topic> [modalities] [language]")
        print("  modalities: text,image,audio,document,code,diagram (comma-separated)")
        print("  language: he, en")
        sys.exit(1)

    topic = sys.argv[1]
    modalities = sys.argv[2].split(",") if len(sys.argv) > 2 else ["text", "diagram"]
    language = sys.argv[3] if len(sys.argv) > 3 else "he"

    engine = get_multimodal_engine()
    result = engine.explain(topic, modalities=modalities, language=language)

    print(result.summary())
    print()
    if result.text:
        print(result.text[:500])
